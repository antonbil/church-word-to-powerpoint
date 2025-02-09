# sermon_extract.py
from docx import Document
import re
from datetime import datetime
import locale
class SermonExtract:
    """
    Contains the methods for extracting information from the Word document.
    """

    def extract_hymn_section(self, paragraphs):
        """
        Extracts the hymn section, including title, and all hymns within the section
        (text and images) from a list of paragraphs.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            tuple: (title, hymn_data)
                   title (str or None): The title of the hymn section (or None if no title).
                   hymn_data (list): A list of dictionaries, where each dictionary represents a
                                     hymn and contains its text and image data.
        """
        hymn_data = []
        current_text = []
        title = None
        new_index = 0
        in_hymn_section = False
        index = -1
        print("check for hymn sections")
        # print(self.current_paragraph_index)

        # check if the hymn-section has a title:
        if len(paragraphs) > 0:
            in_hymn_section, index, title = self.get_hymn_title(index, paragraphs)
        while self.current_paragraph_index + index < self.num_paragraphs-1:
            index = index + 1
            paragraph = self.word_document.paragraphs[self.current_paragraph_index + index]

            if self.tags["hymn"]["begin"] in paragraph.text:
                # a new hymn is started
                in_hymn_section = True
                continue
            if self.check_end_tag("hymn", paragraph):
                # no hymn, but a next hymn can be possible
                in_hymn_section = False
                if len(current_text) > 0:
                    paragraph_data = {"text": "\n".join(current_text), "images": []}
                    hymn_data.append(paragraph_data)
                current_text = []
                new_index = index + 1
                break
            if len(paragraph.text.strip()) == 0 and not in_hymn_section:
                # empty line; skip
                new_index = index
                continue
            if len(paragraph.text.strip()) > 0 and not in_hymn_section:
                # not next hymn, so the hymn-section is finished
                new_index = index
                break
            if in_hymn_section:
                self.extract_paragraph_content(paragraph)
                if len(paragraph.text) > 0:
                    current_text.append(paragraph.text)

                # check if the hymn-section starts with a picture:
                if len(hymn_data) == 0:
                    # if first paragraph in hymn-section contains an image, this is considered as a 'hymn'
                    self.get_hymn_image(hymn_data, paragraph)
        self.current_paragraph_index = self.current_paragraph_index + new_index
        return title, hymn_data

    def extract_offering_section(self, paragraphs):
        """
        Extracts the offering section (text) from a list of paragraphs.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            tuple: (offering_data)
                   offering_data (list): A list of dictionaries, where each dictionary
                                        represents a part of the offering and contains
                                        its text.
        """
        print("extract_offering_section")
        def add_line_function(paragraph, current_text, _):
            cleaned_line = re.sub(r' {5,}', '\n', paragraph.text.strip())

            current_text.append(cleaned_line)
        full_text = self. _extract_section_text("offering", add_line_function)
        offering_goal, bank_account_number = self.extract_bank_account_number(full_text)
        offering_data = {"offering_goal": offering_goal, "bank_account_number": bank_account_number}
        return offering_data

    def _extract_section_text(self, section_name, add_line_function = None, outro_data = None):
        """
        Extracts text from a specified section in the Word document.

        Args:
            section_name (str): The name of the section to extract text from (e.g., "offering").

        Returns:
            list: A list of strings, where each string is a cleaned line of text from the section.
            int: the index of the last processed paragraph.
        """
        current_text = []
        in_section = False
        new_index = 0
        index = -1

        while self.current_paragraph_index + index < self.num_paragraphs - 1:
            index += 1
            paragraph = self.word_document.paragraphs[self.current_paragraph_index + index]

            if self.tags[section_name]["begin"] in paragraph.text:
                # A new section has started
                in_section = True
                t = paragraph.text.split(self.tags[section_name]["begin"])[-1]
                if t and add_line_function:
                    paragraph.text = t
                    add_line_function(paragraph, current_text, outro_data)
                continue

            if self.check_end_tag(section_name, paragraph):
                # End tag found for the section
                in_section = False
                new_index = index
                break

            if len(paragraph.text.strip()) == 0 and not in_section:
                # Empty line outside the section, skip it
                new_index = index
                continue

            if len(paragraph.text.strip()) > 0 and not in_section:
                # Text found outside the section, so the section is finished
                new_index = index
                break

            if in_section:
                if len(paragraph.text) > 0 and add_line_function:
                    add_line_function(paragraph, current_text, outro_data)

                new_index = index
        current_text = "\n".join(current_text)
        self.current_paragraph_index = self.current_paragraph_index + new_index

        return current_text


    def extract_intro_section(self, paragraphs):
        """
        Extracts the introduction section (date, time, parson, theme, organist) from a list of paragraphs.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            dict: A dictionary containing the extracted information:
                  {
                      "date": "Zondag 5 januari 2025",
                      "time": "10.00 uur",
                      "parson": "ds. Elly van Kuijk-Spaans",
                      "theme": "“Zaaien: een gids voor beginners”",
                      "organist": "Martin van der Bent"
                  }
        """
        print("extract_intro_section")
        intro_data = {}
        current_text = []

        def add_line_function(paragraph, current_text, _):
            current_text.append(paragraph.text.strip())
        intro_text = self. _extract_section_text("intro", add_line_function)

        sermon = self.settings.get_setting('word-intro-date_label')
        # Extract date
        sermon_date_list = [l for l in current_text if sermon in l]
        date_text = "25 december 2024"
        if len(sermon_date_list) > 0:
            l2 = sermon_date_list[0].split(sermon)
            if len(l2) > 1:
                date_text = l2[1].strip()

        # Convert month names to numbers
        month_numbers = {
            "januari": "1",
            "februari": "2",
            "maart": "3",
            "april": "4",
            "mei": "5",
            "juni": "6",
            "juli": "7",
            "augustus": "8",
            "september": "9",
            "oktober": "10",
            "november": "11",
            "december": "12",
        }

        intro_data["date"] = date_text
        if date_text:
            day, month_name, year = date_text.split()
            month_number = month_numbers.get(month_name.lower())
            if month_number:
                numeric_date_str = f"{day} {month_number} {year}"
                try:
                    date_object = datetime.strptime(numeric_date_str, "%d %m %Y")
                    intro_data["date"] = f"{self.get_day_of_week(date_object.weekday())} {date_text}"
                except ValueError as e:
                    print(f"Error parsing date: {e}")

        # Extract time
        time_match = re.search(r"(\d{1,2}\.\d{2}\suur)", intro_text)
        if time_match:
            intro_data["time"] = time_match.group(1)

        # Extract parson
        parson_text = self.settings.get_setting("word-intro-parson_text")
        parson_match = re.search(rf"{parson_text}\s*\n\s*(.+)", intro_text)
        if parson_match:
            intro_data["parson"] = parson_match.group(1).strip()

        # Regex to find "Thema:" and capture the text on the following line(s)
        theme_text = self.settings.get_setting("word-intro-theme_text")
        theme_match = re.search(rf"{theme_text}\s*“([^“”]+)”", intro_text)
        if theme_match:
            theme = theme_match.group(1).strip()
            intro_data["theme"] = theme

        # Extract organ-player and performed piece
        organ_text = self.settings.get_setting("word-intro-organplayer_text")
        organist_match = re.search(rf"{organ_text}\s+(.+):\s+(.+)", intro_text)
        if organist_match:
            intro_data["organist"] = organist_match.group(1).strip()
            intro_data["performed_piece"] = organist_match.group(2).strip()

        # print(intro_data)
        return intro_data

    def extract_reading_section(self, paragraphs):
        """
        Extracts the reading section from a list of paragraphs.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            tuple: (title, reading_data)
                   title (str or None): The title of the reading section (or None if no title).
                   reading_data (list): A list of dictionaries, where each dictionary
                                        represents a part of the reading and contains
                                        its text (and potentially images, though they're not expected here).
        """
        print("extract_reading_section")
        reading_data = []
        current_text = []
        title = None
        new_index = 0
        in_reading_section = False
        index = -1
        # check if the reading-section has a title:
        if len(paragraphs) > 0:
            in_reading_section, index, title = self.get_reading_title(index, paragraphs)
        while self.current_paragraph_index + index < self.num_paragraphs-1:
            index = index + 1
            paragraph = self.word_document.paragraphs[self.current_paragraph_index + index]

            if self.tags["reading"]["begin"] in paragraph.text:
                # a new reading is started
                in_reading_section = True
                continue
            if self.check_end_tag("reading", paragraph):
                # no reading, but a next reading can be possible
                in_reading_section = False
                new_index = index + 1
                break
            if len(paragraph.text.strip()) == 0 and not in_reading_section:
                # empty line; skip
                new_index = index
                continue
            if len(paragraph.text.strip()) > 0 and not in_reading_section:
                # not next reading, so the reading-section is finished
                new_index = index
                break
            if in_reading_section:
                if len(paragraph.text) > 0:
                    cleaned_line = re.sub(r' {5,}', '\n', paragraph.text.strip())
                    current_text.append(cleaned_line)
        # Split the text into multiple parts if needed
        full_text = "\n".join(current_text)
        print(full_text)
        # def add_line_function(paragraph, current_text, _):
        #     print(paragraph.text)
        #     cleaned_line = re.sub(r' {5,}', '\n', paragraph.text.strip())
        #     current_text.append(cleaned_line)
        # full_text = self. _extract_section_text("reading", add_line_function)

        lines = full_text.split('\n')

        if len(lines) > self.max_reading_lines:
            parts = [lines[i:i + self.max_reading_lines] for i in range(0, len(lines), self.max_reading_lines)]
            for part in parts:
                reading_data.append({"text": "\n".join(part), "images": []})
        else:
            reading_data.append({"text": full_text, "images": []})
        self.current_paragraph_index = self.current_paragraph_index + new_index
        return title, reading_data

    def extract_illustration(self, paragraphs):
        """
        Extracts the illustration from the given paragraphs.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            tuple: (image_data, image_content_type) or (None, None) if no illustration is found.
        """
        print("extract_illustration")

        in_illustration_section = False
        image = None
        index = -1
        illustration_data = []
        while self.current_paragraph_index + index < self.num_paragraphs-1:
            index = index + 1
            paragraph = self.word_document.paragraphs[self.current_paragraph_index + index]
            if self.tags["illustration"]["begin"] in paragraph.text:
                in_illustration_section = True
                continue
            if self.check_end_tag("illustration", paragraph):
                in_illustration_section = False
                break
            if in_illustration_section:

                    self.extract_paragraph_content(paragraph)

                    # if first paragraph in hymn-section contains an image, this is considered as a 'hymn'
                    self.get_hymn_image(illustration_data, paragraph)

        self.current_paragraph_index = self.current_paragraph_index + index
        if len(illustration_data) > 0:
            image = illustration_data[0]["images"][0]
        return image


    def extract_outro_section(self, paragraphs):
        """
        Extracts the date and parson from the outro section.

        Args:
            paragraphs (list): A list of paragraphs (docx.paragraph.Paragraph objects).

        Returns:
            tuple: (date, parson) or (None, None) if not found.
        """
        print("extract_outro_section")
        current_text = []

        # Use a dictionary to store the values, that will be used as the closure for the function add_line_function
        outro_data = {
            "date_text": "",
            "parson": "",
            "performed_piece": "",
            "previous_line_is_sermon": False,
            "organ_text": self.settings.get_setting("word-outro-organ_text"),
            "next_sermon_text": self.settings.get_setting("word-outro-next_sermon_text")
        }

        # define the function, that will use the closure that is defined in the lines above
        def add_line_function(paragraph, _, outro_data):
            if outro_data["organ_text"] in paragraph.text:
                organ_text = outro_data["organ_text"]
                performed_piece_match = re.search(rf"{organ_text}\s*(.+)", paragraph.text.strip()) # use f-string
                if performed_piece_match:
                    outro_data["performed_piece"] = performed_piece_match.group(1).strip()
            if outro_data["next_sermon_text"].lower() in paragraph.text.lower():
                outro_data["previous_line_is_sermon"] = True
                return

            if outro_data["previous_line_is_sermon"]:
                parts = paragraph.text.split('\t')
                outro_data["previous_line_is_sermon"] = False
                if len(parts) >= 2:
                    date_text1 = parts[0]
                    parson1 = parts[len(parts) - 1]
                    date_text1 = self.format_date(date_text1)
                    outro_data["date_text"] = date_text1
                    outro_data["parson"] = parson1
        self._extract_section_text("outro", add_line_function, outro_data)
        #make sure the program will end here
        self.current_paragraph_index = 200000

        return outro_data["date_text"], outro_data["parson"], outro_data["performed_piece"]

    def extract_bank_account_number(self, text):
        """
        Extracts the offering goal and bank account number from the text.

        Args:
            text (str): The text from the offering section.

        Returns:
            tuple: (offering_goal, bank_account_number)
        """
        bank_account_number = ""
        offering_goal = ""
        # Split the text at the "blauwe zak"
        blue_bag_text = self.settings.get_setting("word-offering-blue_bag_text")
        parts = text.split(blue_bag_text)
        first_part_of_offering_text = parts[0]

        # Regex to find the bank account number (IBAN) in the first part
        bank_account_regex = "([A-Z]{2}\d{2}\s[A-Z]{4}\s\d{4}\s\d{4}\s\d{2})"
        bank_number_match = re.search(bank_account_regex, first_part_of_offering_text)
        if bank_number_match:
            bank_account_number = bank_number_match.group(1).strip()

        # Clean the bank account number
        bank_account_number = re.sub(r"[^a-zA-Z0-9]", " ", bank_account_number)  # Replace non-alphanumeric with space
        bank_account_number = re.sub(r"\s{2,}", " ", bank_account_number)  # Replace multiple spaces with one
        bank_account_number = bank_account_number.strip()

        # Regex to find offering goal (text after "1ste (rode zak) Diaconie:")
        red_bag_text = self.settings.get_setting("word-offering-red_bag_text")
        diaconie_text = self.settings.get_setting("word-offering-diaconie_text")
        # the old regex: r"1ste \(rode zak\) Diaconie:\s*(.*?)(?=\n|$)"
        offering_goal_regex = fr"1ste \({red_bag_text}\)\s*{diaconie_text}\s*(.*?)(?=\n|$)"
        goal_match = re.search(offering_goal_regex, text)
        if goal_match:
            offering_goal = goal_match.group(1).strip()
        return offering_goal, bank_account_number


